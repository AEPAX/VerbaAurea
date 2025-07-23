#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文档处理服务
"""

import os
import time
import sys
import tempfile
from typing import Dict, Any, Tuple
from pathlib import Path

# 添加项目根目录到Python路径以导入现有模块
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

# 使用新架构的适配器
try:
    from verba_aurea.interfaces.api.service_adapter import get_api_adapter
    USE_NEW_ARCHITECTURE = True
except ImportError:
    # 降级到旧架构
    from document_processor import insert_split_markers
    from text_analysis import extract_elements_info
    from config_manager import DEFAULT_CONFIG
    USE_NEW_ARCHITECTURE = False

from ..models.request_models import ProcessingConfig
from ..models.response_models import ProcessingResult
from ..utils.exceptions import DocumentProcessingError
from ..utils.file_handler import FileHandler


class DocumentProcessingService:
    """文档处理服务类"""
    
    def __init__(self):
        self.file_handler = FileHandler()
    
    async def process_document(
        self,
        input_file_path: str,
        config: ProcessingConfig = None,
        debug_mode: bool = False
    ) -> Tuple[str, ProcessingResult]:
        """
        处理文档并返回结果

        Args:
            input_file_path: 输入文件路径
            config: 处理配置
            debug_mode: 是否启用调试模式

        Returns:
            Tuple[str, ProcessingResult]: (输出文件路径, 处理结果)
        """
        start_time = time.time()

        try:
            if USE_NEW_ARCHITECTURE:
                # 使用新架构
                adapter = get_api_adapter()
                config_dict = config.dict() if config else None
                output_file_path, result_dict = await adapter.process_document(
                    input_file_path, config_dict, debug_mode
                )

                # 转换结果格式
                processing_result = ProcessingResult(
                    split_count=result_dict.get('split_count', 0),
                    processing_time=result_dict.get('processing_time', 0.0),
                    file_size_before=result_dict.get('file_size_before', 0),
                    file_size_after=result_dict.get('file_size_after', 0),
                    elements_stats=result_dict.get('elements_stats', {}),
                    success=result_dict.get('success', False),
                    message=result_dict.get('message', ''),
                    error=result_dict.get('error')
                )

                return output_file_path, processing_result

            else:
                # 使用旧架构
                # 准备配置
                processing_config = self._prepare_config(config, debug_mode)

                # 获取输入文件信息
                input_file_size = self.file_handler.get_file_size(input_file_path)

                # 生成输出文件路径
                original_filename = os.path.basename(input_file_path)
                file_id = Path(input_file_path).stem.split('_')[0]  # 提取文件ID
                output_file_path = self.file_handler.create_output_path(file_id, original_filename)

                # 预处理：分析文档结构
                doc_stats = await self._analyze_document(input_file_path, processing_config)

                # 执行文档处理
                success = insert_split_markers(
                    input_file=input_file_path,
                    output_file=output_file_path,
                    config=processing_config
                )

                if not success:
                    raise DocumentProcessingError("文档处理失败")
            
            # 获取输出文件信息
            output_file_size = self.file_handler.get_file_size(output_file_path)
            
            # 计算处理时间
            processing_time = time.time() - start_time
            
            # 分析处理结果
            result_stats = await self._analyze_processing_result(
                output_file_path, processing_config
            )
            
            # 构建处理结果
            processing_result = ProcessingResult(
                split_count=result_stats.get('split_count', 0),
                processing_time=round(processing_time, 4),
                file_size_before=input_file_size,
                file_size_after=output_file_size,
                element_count=doc_stats.get('element_count', 0),
                paragraph_count=doc_stats.get('paragraph_count', 0),
                table_count=doc_stats.get('table_count', 0),
                image_count=doc_stats.get('image_count', 0),
                has_images=doc_stats.get('has_images', False),
                heading_count=doc_stats.get('heading_count', 0),
                split_points=result_stats.get('split_points', []),
                avg_segment_length=result_stats.get('avg_segment_length'),
                min_segment_length=result_stats.get('min_segment_length'),
                max_segment_length=result_stats.get('max_segment_length')
            )
            
            return output_file_path, processing_result
            
        except Exception as e:
            # 清理可能创建的文件
            if 'output_file_path' in locals():
                self.file_handler.cleanup_file(output_file_path)
            
            if isinstance(e, DocumentProcessingError):
                raise
            else:
                raise DocumentProcessingError(
                    f"文档处理过程中发生错误: {str(e)}",
                    details={"error_type": type(e).__name__}
                )
    
    def _prepare_config(self, config: ProcessingConfig = None, debug_mode: bool = False) -> Dict[str, Any]:
        """准备处理配置"""
        if config is None:
            # 使用默认配置
            processing_config = DEFAULT_CONFIG.copy()
        else:
            # 转换为现有系统的配置格式
            processing_config = config.to_legacy_config()
        
        # 覆盖调试模式 - API环境下禁用调试输出避免编码问题
        processing_config['processing_options']['debug_mode'] = False
        
        # 设置临时输出目录
        processing_config['processing_options']['output_folder'] = self.file_handler.output_dir
        
        return processing_config
    
    async def _analyze_document(self, file_path: str, config: Dict[str, Any]) -> Dict[str, Any]:
        """分析文档结构"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取元素信息
            table_factor = config.get("document_settings", {}).get("table_length_factor", 1.0)
            image_factor = config.get("document_settings", {}).get("image_length_factor", 100)
            debug_mode = config.get("processing_options", {}).get("debug_mode", False)
            
            elements_info = extract_elements_info(doc, table_factor, debug_mode, image_factor)
            
            # 统计信息
            paragraph_count = sum(1 for elem in elements_info if elem['type'] == 'para')
            table_count = sum(1 for elem in elements_info if elem['type'] == 'table')
            heading_count = sum(1 for elem in elements_info if elem.get('is_heading', False))
            image_count = sum(elem.get('image_count', 0) for elem in elements_info)
            has_images = any(elem.get('has_images', False) for elem in elements_info)
            
            return {
                'element_count': len(elements_info),
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'heading_count': heading_count,
                'image_count': image_count,
                'has_images': has_images
            }
            
        except Exception as e:
            raise DocumentProcessingError(
                f"文档分析失败: {str(e)}",
                details={"file_path": file_path}
            )
    
    async def _analyze_processing_result(self, output_file_path: str, config: Dict[str, Any]) -> Dict[str, Any]:
        """分析处理结果"""
        try:
            from docx import Document
            
            doc = Document(output_file_path)
            
            # 统计分隔符数量
            split_count = 0
            split_points = []
            segment_lengths = []
            current_length = 0
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip() == "<!--split-->":
                    split_count += 1
                    split_points.append(i)
                    if current_length > 0:
                        segment_lengths.append(current_length)
                    current_length = 0
                else:
                    current_length += len(para.text)
            
            # 添加最后一个段落的长度
            if current_length > 0:
                segment_lengths.append(current_length)
            
            # 计算统计信息
            avg_segment_length = sum(segment_lengths) / len(segment_lengths) if segment_lengths else 0
            min_segment_length = min(segment_lengths) if segment_lengths else 0
            max_segment_length = max(segment_lengths) if segment_lengths else 0
            
            return {
                'split_count': split_count,
                'split_points': split_points,
                'avg_segment_length': round(avg_segment_length, 2),
                'min_segment_length': min_segment_length,
                'max_segment_length': max_segment_length
            }
            
        except Exception as e:
            # 如果分析失败，返回基本信息
            return {
                'split_count': 0,
                'split_points': [],
                'avg_segment_length': None,
                'min_segment_length': None,
                'max_segment_length': None
            }
