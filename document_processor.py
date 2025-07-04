#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文档处理核心模块
处理Word文档，在适当位置插入分隔标记
"""

from docx import Document
from text_analysis import extract_elements_info
import os
from text_analysis import (
    is_sentence_boundary,
    find_nearest_sentence_boundary
)
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import qn
from io import BytesIO
from copy import deepcopy


def insert_split_markers(input_file, output_file, config):
    """
    在Word文档中根据算法自动嵌入<!--split-->标记

    使用多种策略确定分隔点:
    1. 标题识别
    2. 语义结构分析
    3. 长度控制
    4. 自然段落边界
    5. 句子完整性保证
    """
    # 从配置中提取参数
    doc_settings = config["document_settings"]
    proc_options = config["processing_options"]
    adv_settings = config["advanced_settings"]

    max_length = doc_settings["max_length"]
    min_length = doc_settings["min_length"]
    sentence_integrity_weight = doc_settings["sentence_integrity_weight"]
    debug_mode = proc_options["debug_mode"]
    search_window = adv_settings["search_window"]
    min_split_score = adv_settings["min_split_score"]
    heading_score_bonus = adv_settings["heading_score_bonus"]
    sentence_end_score_bonus = adv_settings["sentence_end_score_bonus"]
    length_score_factor = adv_settings["length_score_factor"]

    if debug_mode:
        print(f"正在处理文档: {input_file}")

    # 如果配置为跳过已存在文件，且输出文件已存在，则跳过处理
    if proc_options["skip_existing"] and os.path.exists(output_file):
        if debug_mode:
            print(f"跳过已存在文件: {output_file}")
        return True

    try:
        doc = Document(input_file)
    except Exception as e:
        print(f"无法打开文档 {input_file}: {str(e)}")
        return False

    # 创建新文档
    new_doc = Document()

    # 统一提取段落 + 表格 + 图片
    table_factor = config.get("document_settings", {}).get("table_length_factor", 1.0)
    image_factor = config.get("document_settings", {}).get("image_length_factor", 100)
    elements_info = extract_elements_info(doc, table_factor, debug_mode, image_factor)

    if debug_mode:
        print(f"文档共有 {len(elements_info)} 个元素（段落 + 表格）")

    # 确定分隔点
    split_points = find_split_points(
        elements_info,
        max_length, min_length,
        sentence_integrity_weight, search_window,
        min_split_score, heading_score_bonus,
        sentence_end_score_bonus, length_score_factor,
        debug_mode, adv_settings
    )

    # 后处理：检查所有分割点确保不会打断句子
    final_split_points = refine_split_points(
        elements_info, split_points, search_window, debug_mode
    )

    final_split_points = merge_heading_with_body(elements_info, final_split_points)

    if debug_mode:
        print(f"最终分割点: {final_split_points}")

    # 创建带有分隔符的新文档
    result = create_output_document(
        doc,
        new_doc,
        final_split_points,
        output_file,
        debug_mode,
        config
    )

    return result


# ---------- 重新实现分割三大函数 ----------
def find_split_points(elements_info, max_length, min_length,
                      sentence_integrity_weight, search_window,
                      min_split_score, heading_score_bonus,
                      sentence_end_score_bonus, length_score_factor,
                      debug_mode, adv_settings):
    """
    寻找分割点：
      1. 遇到标题立即插分割点（若开启 force_split_before_heading）。
      2. 标题后的【空行 + 前 cooldown 段正文/表格】全部跳过，不参与打分。
         cooldown 数由 advanced_settings.heading_cooldown_elements 控制，默认 2。
    """
    force_heading = adv_settings.get("force_split_before_heading", True)
    cooldown_len  = adv_settings.get("heading_cooldown_elements", 2)

    split_points = []
    current_length = 0
    last_potential = -1
    cooldown_after_heading = 0   # ← 冷却段计数器

    for idx, elem in enumerate(elements_info):

        # ---------- 标题：强制分段 ----------
        if elem['is_heading'] and idx > 0 and force_heading:
            if not split_points or idx != split_points[-1]:
                split_points.append(idx)
            current_length = 0
            last_potential = idx
            cooldown_after_heading = cooldown_len     # ← 开启冷却
            if debug_mode:
                prev = (elem['text'][:30] + '...') if elem['text'] else '[table]'
                print(f"  #{idx:03d} (heading) 强制分段 «{prev}»")
            continue

        # ---------- 空行：始终累长，绝不当候选 ----------
        if elem['length'] == 0:
            current_length += elem['length']
            continue

        # ---------- 冷却阶段：仅累长，不打分 ----------
        if cooldown_after_heading > 0:
            current_length += elem['length']
            cooldown_after_heading -= 1
            continue

        # ---------- 计算得分 ----------
        current_length += elem['length']
        score = calculate_split_score(
            idx, elem, elements_info, current_length,
            min_length, max_length, sentence_integrity_weight,
            heading_score_bonus, sentence_end_score_bonus,
            length_score_factor, split_points,
            adv_settings
        )

        if debug_mode:
            pv = (elem['text'][:30] + '...') if elem['text'] else '[table]'
            print(f"  #{idx:03d} ({elem['type']}) len={elem['length']} score={score:.1f} {pv}")

        # ---------- 命中分割 ----------
        if score >= min_split_score and idx > 0:
            split_points.append(idx)
            current_length = 0
            last_potential = idx

        # ---------- 超长兜底 ----------
        elif current_length > max_length * 1.5:
            best = find_nearest_sentence_boundary(elements_info, idx, search_window)
            if best >= 0 and (not split_points or best > split_points[-1]):
                split_points.append(best)
                current_length = 0
                last_potential = best
            elif idx - last_potential > 3:
                split_points.append(idx)
                current_length = 0
                last_potential = idx

    return split_points


def calculate_split_score(idx, elem, elements_info, current_length,
                          min_length, max_length, sentence_integrity_weight,
                          heading_score_bonus, sentence_end_score_bonus,
                          length_score_factor, split_points, adv_settings):

    score = 0
    heading_after_penalty = adv_settings.get("heading_after_penalty", 12)

    # 基础分
    if elem['type'] == 'para':
        if elem['is_heading']:
            score += heading_score_bonus
        if elem['ends_with_period']:
            score += sentence_end_score_bonus

        if idx > 0 and elements_info[idx-1]['type'] == 'para' and \
           is_sentence_boundary(elements_info[idx-1]['text'], elem['text']):
            score += sentence_integrity_weight
        else:
            score -= 10
    else:
        score += 6          # 表格基分

    # 紧跟标题统一扣分
    prev = idx - 1
    while prev >= 0 and elements_info[prev]['type'] == 'para' and \
          elements_info[prev]['length'] == 0:
        prev -= 1
    if prev >= 0 and elements_info[prev]['is_heading']:
        score -= heading_after_penalty

    # 长度因子
    if current_length >= min_length:
        score += min(4, (current_length - min_length)//length_score_factor)
    elif current_length < min_length * 0.7:
        score -= 5

    # 距上个分割点太近
    if split_points and idx - split_points[-1] < 3:
        score -= 8

    # 超长补分
    if current_length > max_length:
        score += 4

    return score


def refine_split_points(elements_info, split_points, search_window, debug_mode):
    refined = []

    for sp in split_points:
        # 若分割点本身或紧邻标题，则完全保留
        if elements_info[sp]['is_heading'] \
           or (sp > 0 and elements_info[sp-1]['is_heading']):
            refined.append(sp)
            continue

        # 仅对 “段落 ↔ 段落” 之间尝试句边界微调
        need_adjust = False
        if sp > 0 and \
           elements_info[sp-1]['type'] == 'para' and \
           elements_info[sp]['type'] == 'para':
            need_adjust = not is_sentence_boundary(elements_info[sp-1]['text'],
                                                   elements_info[sp]['text'])

        if need_adjust:
            best = find_nearest_sentence_boundary(elements_info, sp, search_window)
            refined.append(best if best >= 0 else sp)
        else:
            refined.append(sp)

    return sorted(set(refined))

def merge_heading_with_body(elements_info, split_points):
    """
    保证“标题 + 第一块真实内容（段落或表格）”不可被拆开
    —— 无论分割点落在空行/表格/段落，只要位于这两者之间就删除。
    """
    if not split_points:
        return []

    keep = set(split_points)

    for sp in split_points:
        # 寻找 sp 前方最近非空元素
        i = sp - 1
        while i >= 0 and elements_info[i]['length'] == 0:
            i -= 1

        if i >= 0 and elements_info[i]['is_heading']:
            heading_idx = i

            # 找标题后的首块非空内容
            j = heading_idx + 1
            while j < len(elements_info) and elements_info[j]['length'] == 0:
                j += 1

            first_content_idx = j

            if heading_idx < sp <= first_content_idx:
                keep.discard(sp)

    return sorted(keep)





def copy_image_relationships(source_doc, target_doc, debug_mode=False):
    """
    复制源文档中的图片关系到目标文档
    返回一个映射字典：{old_rId: new_rId}
    """
    rId_mapping = {}
    try:
        for rId in source_doc.part.rels:
            rel = source_doc.part.rels[rId]
            if "image" in rel.target_ref:
                try:
                    # 获取图片二进制数据
                    img_part = rel.target_part
                    img_data = img_part._blob

                    # 创建新的图片部分
                    my_bytesio = BytesIO(img_data)
                    image_part = target_doc.part.package.get_or_add_image_part(my_bytesio)

                    # 创建新的关系
                    new_rId = target_doc.part.relate_to(image_part, 'image')
                    rId_mapping[rId] = new_rId

                    if debug_mode:
                        print(f"  图片关系复制: {rId} -> {new_rId} ({len(img_data)/1024:.1f}KB)")

                except Exception as e:
                    if debug_mode:
                        print(f"  警告: 复制图片关系 {rId} 时出错: {str(e)}")
    except Exception as e:
        if debug_mode:
            print(f"  警告: 复制图片关系时出错: {str(e)}")

    return rId_mapping


def copy_single_table(table, new_doc, debug_mode):
    """复制单个表格"""
    if table is None:  # ← 兜底
        return
    try:
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if rows > 0 else 0

        if rows > 0 and cols > 0:
            new_table = new_doc.add_table(rows=rows, cols=cols)

            try:
                new_table.style = table.style
            except:
                pass

            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                        try:
                            new_cell = new_table.rows[i].cells[j]
                            if cell.text:
                                new_cell.text = cell.text
                        except:
                            pass
    except Exception as e:
        if debug_mode:
            print(f"  警告: 处理表格时出错: {str(e)}")


def create_output_document(doc, new_doc, split_points, output_file, debug_mode, config=None):
    split_marker_cnt = 0
    para_iter = iter(doc.paragraphs)
    tbl_iter  = iter(doc.tables)
    next_para = next(para_iter, None)
    next_tbl  = next(tbl_iter, None)
    idx = -1

    # 检查是否需要保留图片
    preserve_images = True
    if config:
        preserve_images = config.get("document_settings", {}).get("preserve_images", True)

    # 复制图片关系（如果启用图片保留）
    rId_mapping = {}
    if preserve_images:
        rId_mapping = copy_image_relationships(doc, new_doc, debug_mode)
        if debug_mode and rId_mapping:
            print(f"  复制了 {len(rId_mapping)} 个图片关系")
    elif debug_mode:
        print("  图片保留功能已禁用")

    # 将 Word DOM 再次顺序遍历
    for el in doc._element.body:
        idx += 1
        if idx in split_points:
            new_doc.add_paragraph("<!--split-->")
            split_marker_cnt += 1

        if isinstance(el, CT_P):
            # —— 段落 ——
            if preserve_images:
                copy_paragraph(next_para, new_doc, debug_mode, rId_mapping)
            else:
                copy_paragraph_text_only(next_para, new_doc, debug_mode)
            next_para = next(para_iter, None)
        elif isinstance(el, CT_Tbl):
            # —— 表格 ——
            copy_single_table(next_tbl, new_doc, debug_mode)
            next_tbl = next(tbl_iter, None)

    # 保存
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    new_doc.save(output_file)
    if debug_mode:
        print(f"✓ 保存: {output_file} (split={split_marker_cnt})")
    return True



def copy_paragraph(src_para, new_doc, debug_mode, rId_mapping=None):
    """
    复制段落内容和格式，包括图片
    rId_mapping: 图片关系ID映射字典
    """
    # 检查段落是否包含图片
    has_images = False
    try:
        for run in src_para.runs:
            if hasattr(run._element, 'xpath'):
                inline_shapes = run._element.xpath('.//w:drawing//wp:inline')
                if inline_shapes:
                    has_images = True
                    break
    except:
        pass

    if has_images and rId_mapping:
        # 如果段落包含图片，需要特殊处理
        copy_paragraph_with_images(src_para, new_doc, debug_mode, rId_mapping)
    else:
        # 原有的文本复制逻辑
        copy_paragraph_text_only(src_para, new_doc, debug_mode)


def copy_paragraph_text_only(src_para, new_doc, debug_mode):
    """复制段落的文本内容和格式（不包括图片）"""
    text = src_para.text
    new_para = new_doc.add_paragraph(text)

    # 复制格式
    try:
        if src_para.style:
            new_para.style = src_para.style
        new_para.alignment = src_para.alignment

        # 复制段落内的文本格式
        for j in range(min(len(src_para.runs), len(new_para.runs))):
            src_run = src_para.runs[j]
            dst_run = new_para.runs[j]

            for attr in ['bold', 'italic', 'underline']:
                if hasattr(src_run, attr):
                    setattr(dst_run, attr, getattr(src_run, attr))

            # 复制字体属性
            if hasattr(src_run, 'font') and hasattr(dst_run, 'font'):
                for attr in ['size', 'name', 'color']:
                    try:
                        if hasattr(src_run.font, attr):
                            src_value = getattr(src_run.font, attr)
                            if src_value:
                                setattr(dst_run.font, attr, src_value)
                    except:
                        pass
    except Exception as e:
        if debug_mode:
            print(f"  警告: 复制格式时出错: {str(e)}")


def copy_paragraph_with_images(src_para, new_doc, debug_mode, rId_mapping):
    """复制包含图片的段落"""
    try:
        # 创建新段落
        new_para = new_doc.add_paragraph()

        # 复制段落级别的格式
        if src_para.style:
            new_para.style = src_para.style
        new_para.alignment = src_para.alignment

        # 逐个复制runs，包括文本和图片
        for src_run in src_para.runs:
            # 检查run是否包含图片
            has_inline_images = False
            try:
                if hasattr(src_run._element, 'xpath'):
                    inline_shapes = src_run._element.xpath('.//w:drawing//wp:inline')
                    has_inline_images = len(inline_shapes) > 0
            except:
                pass

            if has_inline_images:
                # 复制包含图片的run
                copy_run_with_images(src_run, new_para, debug_mode, rId_mapping)
            else:
                # 复制普通文本run
                copy_text_run(src_run, new_para, debug_mode)

    except Exception as e:
        if debug_mode:
            print(f"  警告: 复制包含图片的段落时出错: {str(e)}")
        # 降级到纯文本复制
        copy_paragraph_text_only(src_para, new_doc, debug_mode)


def copy_text_run(src_run, new_para, debug_mode):
    """复制文本run"""
    try:
        new_run = new_para.add_run(src_run.text)

        # 复制run级别的格式
        for attr in ['bold', 'italic', 'underline']:
            if hasattr(src_run, attr):
                setattr(new_run, attr, getattr(src_run, attr))

        # 复制字体属性
        if hasattr(src_run, 'font') and hasattr(new_run, 'font'):
            for attr in ['size', 'name', 'color']:
                try:
                    if hasattr(src_run.font, attr):
                        src_value = getattr(src_run.font, attr)
                        if src_value:
                            setattr(new_run.font, attr, src_value)
                except:
                    pass
    except Exception as e:
        if debug_mode:
            print(f"  警告: 复制文本run时出错: {str(e)}")


def copy_run_with_images(src_run, new_para, debug_mode, rId_mapping):
    """复制包含图片的run"""
    try:
        # 先添加文本部分
        if src_run.text:
            text_run = new_para.add_run(src_run.text)
            # 复制文本格式
            for attr in ['bold', 'italic', 'underline']:
                if hasattr(src_run, attr):
                    setattr(text_run, attr, getattr(src_run, attr))

        # 处理图片
        if hasattr(src_run._element, 'xpath'):
            inline_shapes = src_run._element.xpath('.//w:drawing//wp:inline')
            for shape in inline_shapes:
                try:
                    copy_inline_image(shape, new_para, debug_mode, rId_mapping)
                except Exception as e:
                    if debug_mode:
                        print(f"  警告: 复制内联图片时出错: {str(e)}")

    except Exception as e:
        if debug_mode:
            print(f"  警告: 复制包含图片的run时出错: {str(e)}")


def copy_inline_image(shape_element, new_para, debug_mode, rId_mapping):
    """复制内联图片"""
    try:
        # 获取图片的关系ID
        blip = shape_element.xpath('.//a:blip')[0] if shape_element.xpath('.//a:blip') else None
        if blip is None:
            return

        old_rId = blip.get(qn('r:embed'))
        if old_rId not in rId_mapping:
            if debug_mode:
                print(f"  警告: 找不到图片关系映射: {old_rId}")
            return

        new_rId = rId_mapping[old_rId]

        # 复制整个drawing元素并更新关系ID
        drawing_element = shape_element.getparent().getparent()  # 获取w:drawing元素
        new_drawing = deepcopy(drawing_element)

        # 更新新drawing中的关系ID
        for new_blip in new_drawing.xpath('.//a:blip'):
            if new_blip.get(qn('r:embed')) == old_rId:
                new_blip.set(qn('r:embed'), new_rId)

        # 将新的drawing添加到新段落的run中
        new_run = new_para.add_run()
        new_run._element.append(new_drawing)

        if debug_mode:
            print(f"  图片复制成功: {old_rId} -> {new_rId}")

    except Exception as e:
        if debug_mode:
            print(f"  警告: 复制内联图片时出错: {str(e)}")

