#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
DOCX文档处理器

实现Word文档的处理功能，包括元素提取、分割点计算和文档生成。
"""

import os
import time
from pathlib import Path
from typing import List, Dict, Any
from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from io import BytesIO
from copy import deepcopy

from .base import BaseProcessor
from ..models.document import Document, ProcessingResult, ElementInfo
from ..models.config import ProcessingConfig
from ..analyzers.structure_analyzer import StructureAnalyzer
from ..splitters.semantic_splitter import SemanticSplitter


class DocxProcessor(BaseProcessor):
    """DOCX文档处理器"""
    
    def __init__(self):
        """初始化DOCX处理器"""
        self.structure_analyzer = None
        self.splitter = None
    
    @property
    def supported_extensions(self) -> List[str]:
        """返回支持的文件扩展名列表"""
        return ['.docx', '.doc']
    
    @property
    def processor_name(self) -> str:
        """返回处理器名称"""
        return "DocxProcessor"
    
    def can_process(self, file_path: Path) -> bool:
        """检查是否可以处理指定文件"""
        return file_path.suffix.lower() in self.supported_extensions
    
    def extract_elements(self, file_path: Path, config: ProcessingConfig) -> Document:
        """从文件中提取元素信息"""
        self._initialize_components(config)
        return self.structure_analyzer.analyze_document(file_path)
    
    def process_document(
        self, 
        input_file: Path, 
        output_file: Path, 
        config: ProcessingConfig
    ) -> ProcessingResult:
        """处理文档，插入分隔符并保存"""
        start_time = time.time()
        
        try:
            # 初始化组件
            self._initialize_components(config)
            
            # 获取输入文件大小
            file_size_before = input_file.stat().st_size
            
            # 分析文档结构
            document = self.extract_elements(input_file, config)
            
            if config.processing_options.debug_mode:
                print(f"文档共有 {document.total_elements} 个元素")
                print(f"段落: {document.paragraph_count}, 表格: {document.table_count}, 图片: {document.image_count}")
            
            # 找到分割点
            split_points = self.splitter.find_split_points(document)
            
            if config.processing_options.debug_mode:
                print(f"找到 {len(split_points)} 个分割点: {split_points}")
            
            # 创建输出文档
            success = self._create_output_document(
                input_file, output_file, split_points, config
            )
            
            if not success:
                raise Exception("创建输出文档失败")
            
            # 获取输出文件大小
            file_size_after = output_file.stat().st_size if output_file.exists() else 0
            
            # 计算处理时间
            processing_time = time.time() - start_time
            
            # 统计元素信息
            elements_stats = {
                'paragraphs': document.paragraph_count,
                'tables': document.table_count,
                'images': document.image_count,
                'total_elements': document.total_elements
            }
            
            return ProcessingResult(
                success=True,
                message="文档处理成功",
                split_count=len(split_points),
                processing_time=processing_time,
                file_size_before=file_size_before,
                file_size_after=file_size_after,
                elements_stats=elements_stats,
                split_points=split_points
            )
            
        except Exception as e:
            processing_time = time.time() - start_time
            return ProcessingResult(
                success=False,
                message=f"文档处理失败: {str(e)}",
                processing_time=processing_time,
                error_details=str(e)
            )
    
    def _initialize_components(self, config: ProcessingConfig):
        """初始化处理组件"""
        if self.structure_analyzer is None:
            self.structure_analyzer = StructureAnalyzer(config)
        if self.splitter is None:
            self.splitter = SemanticSplitter(config)
    
    def _create_output_document(
        self, 
        input_file: Path, 
        output_file: Path, 
        split_points: List[int], 
        config: ProcessingConfig
    ) -> bool:
        """创建带有分隔符的输出文档"""
        try:
            # 打开源文档
            source_doc = DocxDocument(input_file)
            
            # 创建新文档
            target_doc = DocxDocument()
            
            # 复制图片关系（如果启用）
            rId_mapping = {}
            if config.document_settings.preserve_images:
                rId_mapping = self._copy_image_relationships(
                    source_doc, target_doc, config.processing_options.debug_mode
                )
            
            # 按顺序处理文档元素
            self._process_document_elements(
                source_doc, target_doc, split_points, config, rId_mapping
            )
            
            # 保存文档
            os.makedirs(output_file.parent, exist_ok=True)
            target_doc.save(output_file)
            
            if config.processing_options.debug_mode:
                print(f"✓ 保存: {output_file} (分割点: {len(split_points)})")
            
            return True
            
        except Exception as e:
            if config.processing_options.debug_mode:
                print(f"创建输出文档时出错: {str(e)}")
            return False
    
    def _copy_image_relationships(
        self, 
        source_doc: DocxDocument, 
        target_doc: DocxDocument, 
        debug_mode: bool
    ) -> Dict[str, str]:
        """复制图片关系"""
        rId_mapping = {}
        try:
            for rId in source_doc.part.rels:
                rel = source_doc.part.rels[rId]
                if "image" in rel.target_ref:
                    try:
                        # 获取图片数据
                        img_part = rel.target_part
                        img_data = img_part._blob
                        
                        # 创建新的图片部分
                        img_stream = BytesIO(img_data)
                        image_part = target_doc.part.package.get_or_add_image_part(img_stream)
                        
                        # 创建新关系
                        new_rId = target_doc.part.relate_to(image_part, 'image')
                        rId_mapping[rId] = new_rId
                        
                        if debug_mode:
                            print(f"  图片关系复制: {rId} -> {new_rId} ({len(img_data)/1024:.1f}KB)")
                    
                    except Exception as e:
                        if debug_mode:
                            print(f"  警告: 复制图片关系 {rId} 时出错: {str(e)}")
        
        except Exception as e:
            if debug_mode:
                print(f"  警告: 复制图片关系时出错: {str(e)}")
        
        return rId_mapping
    
    def _process_document_elements(
        self,
        source_doc: DocxDocument,
        target_doc: DocxDocument,
        split_points: List[int],
        config: ProcessingConfig,
        rId_mapping: Dict[str, str]
    ):
        """按顺序处理文档元素"""
        para_iter = iter(source_doc.paragraphs)
        table_iter = iter(source_doc.tables)
        next_para = next(para_iter, None)
        next_table = next(table_iter, None)
        
        element_index = -1
        
        for element in source_doc._element.body:
            element_index += 1
            
            # 在分割点插入分隔符
            if element_index in split_points:
                target_doc.add_paragraph("<!--split-->")
            
            if isinstance(element, CT_P):
                # 处理段落
                if next_para is not None:
                    if config.document_settings.preserve_images and rId_mapping:
                        self._copy_paragraph_with_images(
                            next_para, target_doc, config.processing_options.debug_mode, rId_mapping
                        )
                    else:
                        self._copy_paragraph_text_only(
                            next_para, target_doc, config.processing_options.debug_mode
                        )
                    next_para = next(para_iter, None)
            
            elif isinstance(element, CT_Tbl):
                # 处理表格
                if next_table is not None:
                    self._copy_table(
                        next_table, target_doc, config.processing_options.debug_mode
                    )
                    next_table = next(table_iter, None)

    def _copy_paragraph_text_only(self, src_para, target_doc, debug_mode: bool):
        """复制段落的文本内容和格式（不包括图片）"""
        try:
            text = src_para.text
            new_para = target_doc.add_paragraph(text)

            # 复制段落格式
            if src_para.style:
                new_para.style = src_para.style
            new_para.alignment = src_para.alignment

            # 复制文本格式
            for i in range(min(len(src_para.runs), len(new_para.runs))):
                src_run = src_para.runs[i]
                dst_run = new_para.runs[i]

                # 复制基本格式
                for attr in ['bold', 'italic', 'underline']:
                    if hasattr(src_run, attr):
                        setattr(dst_run, attr, getattr(src_run, attr))

                # 复制字体属性
                if hasattr(src_run, 'font') and hasattr(dst_run, 'font'):
                    for attr in ['size', 'name', 'color']:
                        try:
                            if hasattr(src_run.font, attr):
                                src_value = getattr(src_run.font, attr)
                                if src_value:
                                    setattr(dst_run.font, attr, src_value)
                        except:
                            pass

        except Exception as e:
            if debug_mode:
                print(f"  警告: 复制段落时出错: {str(e)}")

    def _copy_paragraph_with_images(self, src_para, target_doc, debug_mode: bool, rId_mapping: Dict[str, str]):
        """复制包含图片的段落"""
        try:
            # 检查段落是否包含图片
            has_images = self._paragraph_has_images(src_para)

            if has_images:
                # 创建新段落
                new_para = target_doc.add_paragraph()

                # 复制段落格式
                if src_para.style:
                    new_para.style = src_para.style
                new_para.alignment = src_para.alignment

                # 逐个复制runs
                for src_run in src_para.runs:
                    if self._run_has_images(src_run):
                        self._copy_run_with_images(src_run, new_para, debug_mode, rId_mapping)
                    else:
                        self._copy_text_run(src_run, new_para, debug_mode)
            else:
                # 降级到文本复制
                self._copy_paragraph_text_only(src_para, target_doc, debug_mode)

        except Exception as e:
            if debug_mode:
                print(f"  警告: 复制包含图片的段落时出错: {str(e)}")
            # 降级到文本复制
            self._copy_paragraph_text_only(src_para, target_doc, debug_mode)

    def _copy_table(self, src_table, target_doc, debug_mode: bool):
        """复制表格"""
        try:
            if src_table is None:
                return

            rows = len(src_table.rows)
            cols = len(src_table.rows[0].cells) if rows > 0 else 0

            if rows > 0 and cols > 0:
                new_table = target_doc.add_table(rows=rows, cols=cols)

                # 复制表格样式
                try:
                    new_table.style = src_table.style
                except:
                    pass

                # 复制单元格内容
                for i, row in enumerate(src_table.rows):
                    for j, cell in enumerate(row.cells):
                        if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                            try:
                                new_cell = new_table.rows[i].cells[j]
                                if cell.text:
                                    new_cell.text = cell.text
                            except:
                                pass

        except Exception as e:
            if debug_mode:
                print(f"  警告: 复制表格时出错: {str(e)}")

    def _paragraph_has_images(self, paragraph) -> bool:
        """检查段落是否包含图片"""
        try:
            for run in paragraph.runs:
                if hasattr(run._element, 'xpath'):
                    inline_shapes = run._element.xpath('.//w:drawing//wp:inline')
                    if inline_shapes:
                        return True
            return False
        except:
            return False

    def _run_has_images(self, run) -> bool:
        """检查run是否包含图片"""
        try:
            if hasattr(run._element, 'xpath'):
                inline_shapes = run._element.xpath('.//w:drawing//wp:inline')
                return len(inline_shapes) > 0
            return False
        except:
            return False

    def _copy_text_run(self, src_run, new_para, debug_mode: bool):
        """复制文本run"""
        try:
            new_run = new_para.add_run(src_run.text)

            # 复制格式
            for attr in ['bold', 'italic', 'underline']:
                if hasattr(src_run, attr):
                    setattr(new_run, attr, getattr(src_run, attr))

            # 复制字体属性
            if hasattr(src_run, 'font') and hasattr(new_run, 'font'):
                for attr in ['size', 'name', 'color']:
                    try:
                        if hasattr(src_run.font, attr):
                            src_value = getattr(src_run.font, attr)
                            if src_value:
                                setattr(new_run.font, attr, src_value)
                    except:
                        pass

        except Exception as e:
            if debug_mode:
                print(f"  警告: 复制文本run时出错: {str(e)}")

    def _copy_run_with_images(self, src_run, new_para, debug_mode: bool, rId_mapping: Dict[str, str]):
        """复制包含图片的run"""
        try:
            # 先复制文本部分
            if src_run.text:
                self._copy_text_run(src_run, new_para, debug_mode)

            # 然后处理图片
            if hasattr(src_run._element, 'xpath'):
                inline_shapes = src_run._element.xpath('.//w:drawing//wp:inline')
                for shape in inline_shapes:
                    self._copy_inline_image(shape, new_para, debug_mode, rId_mapping)

        except Exception as e:
            if debug_mode:
                print(f"  警告: 复制包含图片的run时出错: {str(e)}")

    def _copy_inline_image(self, shape_element, new_para, debug_mode: bool, rId_mapping: Dict[str, str]):
        """复制内联图片"""
        try:
            from docx.oxml.ns import qn

            # 获取图片的关系ID
            blip = shape_element.xpath('.//a:blip')[0] if shape_element.xpath('.//a:blip') else None
            if blip is None:
                return

            old_rId = blip.get(qn('r:embed'))
            if old_rId not in rId_mapping:
                if debug_mode:
                    print(f"  警告: 找不到图片关系映射: {old_rId}")
                return

            new_rId = rId_mapping[old_rId]

            # 复制drawing元素并更新关系ID
            drawing_element = shape_element.getparent().getparent()
            new_drawing = deepcopy(drawing_element)

            # 更新关系ID
            for new_blip in new_drawing.xpath('.//a:blip'):
                if new_blip.get(qn('r:embed')) == old_rId:
                    new_blip.set(qn('r:embed'), new_rId)

            # 将图片添加到新段落
            new_run = new_para.add_run()
            new_run._element.append(new_drawing)

            if debug_mode:
                print(f"  复制图片: {old_rId} -> {new_rId}")

        except Exception as e:
            if debug_mode:
                print(f"  警告: 复制图片时出错: {str(e)}")
