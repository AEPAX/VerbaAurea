#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文档结构分析器

负责分析文档的结构，提取段落、表格、图片等元素信息。
"""

from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import qn
from typing import List, Dict, Any, Optional
from pathlib import Path

from ..models.document import Document, ElementInfo
from ..models.config import ProcessingConfig
from .text_analyzer import TextAnalyzer


class StructureAnalyzer:
    """文档结构分析器"""
    
    def __init__(self, config: Optional[ProcessingConfig] = None):
        """初始化结构分析器"""
        self.config = config
        self.text_analyzer = TextAnalyzer(config)
    
    def analyze_document(self, file_path: Path) -> Document:
        """分析文档结构，返回文档对象"""
        try:
            docx_doc = DocxDocument(file_path)
            elements = self.extract_elements_info(docx_doc)
            
            return Document(
                file_path=file_path,
                elements=elements,
                metadata=self._extract_metadata(docx_doc)
            )
        except Exception as e:
            raise ValueError(f"无法分析文档 {file_path}: {str(e)}")
    
    def extract_elements_info(self, docx_doc: DocxDocument) -> List[ElementInfo]:
        """按文档布局顺序提取段落/表格等元素信息"""
        elements = []
        para_idx = -1
        table_idx = -1
        
        # 创建段落和表格的映射
        paragraph_map = {p._element: p for p in docx_doc.paragraphs}
        table_map = {t._element: t for t in docx_doc.tables}
        
        # 按文档顺序遍历所有元素
        for element in docx_doc._element.body:
            if isinstance(element, CT_P):
                # 处理段落
                para_idx += 1
                paragraph = paragraph_map[element]
                element_info = self._analyze_paragraph(paragraph, para_idx)
                elements.append(element_info)
                
            elif isinstance(element, CT_Tbl):
                # 处理表格
                table_idx += 1
                table = table_map[element]
                element_info = self._analyze_table(table, table_idx)
                elements.append(element_info)
        
        return elements
    
    def _analyze_paragraph(self, paragraph, index: int) -> ElementInfo:
        """分析段落元素"""
        text = paragraph.text.strip()
        
        # 分析文本属性
        text_props = self.text_analyzer.analyze_text_properties(text)
        
        # 检查样式是否为标题
        style_is_heading = paragraph.style.name.startswith(('Heading', '标题'))
        is_heading = style_is_heading or text_props['is_heading']
        
        # 检查图片信息
        has_images = self._has_inline_images(paragraph)
        image_count = self._count_inline_images(paragraph)
        image_info = self._get_paragraph_image_info(paragraph) if has_images else []
        
        # 计算长度（包含图片权重）
        base_length = len(text)
        image_length = 0
        if self.config and has_images:
            image_length = image_count * self.config.document_settings.image_length_factor
        total_length = base_length + image_length
        
        return ElementInfo(
            type='para',
            index=index,
            text=text,
            length=total_length,
            base_text_length=base_length,
            is_heading=is_heading,
            is_list_item=text_props['is_list_item'],
            ends_with_period=text_props['ends_with_period'],
            has_images=has_images,
            image_count=image_count,
            image_info=image_info
        )
    
    def _analyze_table(self, table, index: int) -> ElementInfo:
        """分析表格元素"""
        texts = []
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text.strip())
        
        table_text = ' '.join(texts)
        base_length = len(table_text)
        
        # 应用表格长度因子
        total_length = base_length
        if self.config:
            total_length = int(base_length * self.config.document_settings.table_length_factor)
        
        # 收集表格信息
        table_info = {
            'rows': len(table.rows),
            'columns': len(table.columns) if table.rows else 0,
            'cell_count': sum(len(row.cells) for row in table.rows)
        }
        
        return ElementInfo(
            type='table',
            index=index,
            text=table_text,
            length=total_length,
            base_text_length=base_length,
            is_heading=False,
            is_list_item=False,
            ends_with_period=True,  # 表格通常视为完整单元
            has_images=False,
            image_count=0,
            image_info=[],
            table_info=table_info
        )
    
    def _has_inline_images(self, paragraph) -> bool:
        """检查段落是否包含内联图片"""
        try:
            for run in paragraph.runs:
                if hasattr(run._element, 'xpath'):
                    inline_shapes = run._element.xpath('.//w:drawing//wp:inline')
                    if inline_shapes:
                        return True
            return False
        except Exception:
            return False
    
    def _count_inline_images(self, paragraph) -> int:
        """计算段落中内联图片的数量"""
        try:
            count = 0
            for run in paragraph.runs:
                if hasattr(run._element, 'xpath'):
                    inline_shapes = run._element.xpath('.//w:drawing//wp:inline')
                    count += len(inline_shapes)
            return count
        except Exception:
            return 0
    
    def _get_paragraph_image_info(self, paragraph) -> List[Dict[str, Any]]:
        """获取段落中图片的详细信息"""
        try:
            images = []
            for run in paragraph.runs:
                if hasattr(run._element, 'xpath'):
                    inline_shapes = run._element.xpath('.//w:drawing//wp:inline')
                    for shape in inline_shapes:
                        try:
                            # 获取图片尺寸信息
                            extent = shape.xpath('.//wp:extent')[0] if shape.xpath('.//wp:extent') else None
                            width = int(extent.get('cx')) if extent is not None else 0
                            height = int(extent.get('cy')) if extent is not None else 0
                            
                            # 获取图片关系ID
                            blip = shape.xpath('.//a:blip')[0] if shape.xpath('.//a:blip') else None
                            r_embed = blip.get(qn('r:embed')) if blip is not None else None
                            
                            images.append({
                                'width': width,
                                'height': height,
                                'r_embed': r_embed,
                                'element': shape
                            })
                        except Exception:
                            # 如果获取详细信息失败，至少记录存在图片
                            images.append({
                                'width': 0,
                                'height': 0,
                                'r_embed': None,
                                'element': shape
                            })
            return images
        except Exception:
            return []
    
    def _extract_metadata(self, docx_doc: DocxDocument) -> Dict[str, Any]:
        """提取文档元数据"""
        try:
            core_props = docx_doc.core_properties
            return {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created,
                'modified': core_props.modified,
                'revision': core_props.revision
            }
        except Exception:
            return {}
